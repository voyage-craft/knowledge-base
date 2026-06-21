"""Tests for document export: DOCX, HTML, Markdown, LaTeX formats."""

import pytest
import io


async def _login(client, username="admin", password="test-admin-password-123"):
    await client.get("/api/health")
    resp = await client.post("/api/auth/login", json={
        "username": username, "password": password,
    })
    return {"Authorization": f"Bearer {resp.json()['access_token']}"}


def _sample_tiptap():
    return {
        "type": "doc",
        "content": [
            {"type": "heading", "attrs": {"level": 1}, "content": [
                {"type": "text", "text": "Test Heading"},
            ]},
            {"type": "paragraph", "content": [
                {"type": "text", "text": "Hello "},
                {"type": "text", "text": "bold", "marks": [{"type": "bold"}]},
                {"type": "text", "text": " world."},
            ]},
            {"type": "bulletList", "content": [
                {"type": "listItem", "content": [
                    {"type": "paragraph", "content": [{"type": "text", "text": "Item 1"}]},
                ]},
                {"type": "listItem", "content": [
                    {"type": "paragraph", "content": [{"type": "text", "text": "Item 2"}]},
                ]},
            ]},
            {"type": "codeBlock", "content": [
                {"type": "text", "text": "print('hello')"},
            ]},
            {"type": "paragraph", "content": [
                {"type": "text", "text": "A link", "marks": [
                    {"type": "link", "attrs": {"href": "https://example.com"}},
                ]},
            ]},
        ],
    }


async def _create_doc_with_content(client, headers, title="Export Test"):
    resp = await client.post("/api/documents", json={
        "title": title,
        "content_json": _sample_tiptap(),
    }, headers=headers)
    return resp.json()["id"]


# ── DOCX Export ─────────────────────────────────────────────────────────

class TestDOCXExport:

    @pytest.mark.asyncio
    async def test_export_docx_returns_valid_file(self, async_client):
        headers = await _login(async_client)
        doc_id = await _create_doc_with_content(async_client, headers)

        resp = await async_client.get(f"/api/documents/{doc_id}/export?format=docx", headers=headers)
        assert resp.status_code == 200
        assert "application/vnd.openxmlformats-officedocument.wordprocessingml.document" in resp.headers["content-type"]

        # Verify it's a valid DOCX by loading with python-docx
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(resp.content))
        paragraphs = [p.text for p in doc.paragraphs]
        full_text = "\n".join(paragraphs)
        assert "Test Heading" in full_text
        assert "bold" in full_text

    @pytest.mark.asyncio
    async def test_export_docx_contains_lists(self, async_client):
        headers = await _login(async_client)
        doc_id = await _create_doc_with_content(async_client, headers)

        resp = await async_client.get(f"/api/documents/{doc_id}/export?format=docx", headers=headers)
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(resp.content))
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Item 1" in full_text
        assert "Item 2" in full_text

    @pytest.mark.asyncio
    async def test_export_docx_nonexistent_document(self, async_client):
        headers = await _login(async_client)
        resp = await async_client.get("/api/documents/99999/export?format=docx", headers=headers)
        assert resp.status_code == 404


# ── HTML Export ─────────────────────────────────────────────────────────

class TestHTMLExport:

    @pytest.mark.asyncio
    async def test_export_html_contains_content(self, async_client):
        headers = await _login(async_client)
        doc_id = await _create_doc_with_content(async_client, headers)

        resp = await async_client.get(f"/api/documents/{doc_id}/export?format=html", headers=headers)
        assert resp.status_code == 200
        assert "text/html" in resp.headers["content-type"]
        html = resp.content.decode("utf-8")
        assert "<!DOCTYPE html>" in html
        assert "Test Heading" in html
        assert "<strong>bold</strong>" in html

    @pytest.mark.asyncio
    async def test_export_html_xss_prevention(self, async_client):
        """Verify that script tags in content are escaped."""
        headers = await _login(async_client)
        resp = await async_client.post("/api/documents", json={
            "title": '<script>alert("xss")</script>',
            "content_json": {"type": "doc", "content": [
                {"type": "paragraph", "content": [
                    {"type": "text", "text": '<img onerror="alert(1)">'},
                ]},
            ]},
        }, headers=headers)
        doc_id = resp.json()["id"]

        html_resp = await async_client.get(f"/api/documents/{doc_id}/export?format=html", headers=headers)
        html = html_resp.content.decode("utf-8")
        # Script tags should be escaped, not raw
        assert "<script>" not in html


# ── Markdown Export ─────────────────────────────────────────────────────

class TestMarkdownExport:

    @pytest.mark.asyncio
    async def test_export_markdown(self, async_client):
        headers = await _login(async_client)
        doc_id = await _create_doc_with_content(async_client, headers)

        resp = await async_client.get(f"/api/documents/{doc_id}/export?format=markdown", headers=headers)
        assert resp.status_code == 200
        assert "text/markdown" in resp.headers["content-type"]
        md = resp.content.decode("utf-8")
        assert "# Test Heading" in md
        assert "**bold**" in md

    @pytest.mark.asyncio
    async def test_export_markdown_code_block(self, async_client):
        headers = await _login(async_client)
        doc_id = await _create_doc_with_content(async_client, headers)

        resp = await async_client.get(f"/api/documents/{doc_id}/export?format=markdown", headers=headers)
        md = resp.content.decode("utf-8")
        assert "```" in md
        assert "print('hello')" in md


# ── LaTeX Export ────────────────────────────────────────────────────────

class TestLaTeXExport:

    @pytest.mark.asyncio
    async def test_export_latex(self, async_client):
        headers = await _login(async_client)
        doc_id = await _create_doc_with_content(async_client, headers)

        resp = await async_client.get(f"/api/documents/{doc_id}/export?format=latex", headers=headers)
        assert resp.status_code == 200
        assert "application/x-tex" in resp.headers["content-type"]
        tex = resp.content.decode("utf-8")
        assert "\\documentclass" in tex
        assert "\\begin{document}" in tex
        assert "\\section" in tex or "\\subsection" in tex

    @pytest.mark.asyncio
    async def test_export_latex_escapes_special_chars(self, async_client):
        headers = await _login(async_client)
        resp = await async_client.post("/api/documents", json={
            "title": "Special & Chars $100%",
            "content_json": {"type": "doc", "content": [
                {"type": "paragraph", "content": [{"type": "text", "text": "Cost: $50 & tax #1"}]},
            ]},
        }, headers=headers)
        doc_id = resp.json()["id"]

        tex_resp = await async_client.get(f"/api/documents/{doc_id}/export?format=latex", headers=headers)
        tex = tex_resp.content.decode("utf-8")
        # Special chars should be escaped
        assert "\\$" in tex or "\\&" in tex or "\\#" in tex


# ── Invalid Format ──────────────────────────────────────────────────────

class TestInvalidExportFormat:

    @pytest.mark.asyncio
    async def test_invalid_format_rejected(self, async_client):
        headers = await _login(async_client)
        doc_id = await _create_doc_with_content(async_client, headers)
        resp = await async_client.get(f"/api/documents/{doc_id}/export?format=invalid", headers=headers)
        assert resp.status_code == 422
