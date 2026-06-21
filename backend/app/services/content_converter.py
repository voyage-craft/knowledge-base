"""Content format conversion utilities for TipTap JSON documents.

This module provides conversion functions between TipTap JSON and various
export formats (HTML, LaTeX, DOCX, Markdown). It also includes utilities
for text extraction and sanitization.
"""

import json
import re


def sanitize_href(href: str) -> str:
    """Only allow safe URL schemes in link href attributes."""
    href = href.strip()
    if href.lower().startswith(("http://", "https://", "mailto:", "#", "/")):
        return href
    return "#"


def sanitize_filename(name: str) -> str:
    """Sanitize a filename for Content-Disposition header (RFC 5987)."""
    # Remove characters that could cause header injection
    safe = re.sub(r'["\r\n;\\]', '', name)
    # Limit length
    return safe[:200] or "document"


def normalize_content_json(content_json) -> dict | None:
    """Ensure content_json is a dict, handling old double-serialized strings."""
    if content_json is None:
        return None
    if isinstance(content_json, str):
        try:
            content_json = json.loads(content_json)
        except (json.JSONDecodeError, TypeError):
            return None
    return content_json if isinstance(content_json, dict) else None


def _html_escape(s: str) -> str:
    """Escape HTML special characters to prevent XSS."""
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;").replace("'", "&#x27;")


def tiptap_to_html(content_json: dict | None) -> str:
    """Convert TipTap JSON document to HTML string."""
    if not content_json or content_json.get("type") != "doc":
        return ""

    def render_node(node: dict) -> str:
        node_type = node.get("type", "")
        children = "".join(render_node(c) for c in node.get("content", []))

        if node_type == "paragraph":
            if not children:
                return "<p><br></p>"
            return f"<p>{children}</p>"
        elif node_type == "heading":
            level = node.get("attrs", {}).get("level", 1)
            return f"<h{level}>{children}</h{level}>"
        elif node_type == "text":
            text = (node.get("text") or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            marks = node.get("marks", [])
            for mark in marks:
                if mark["type"] == "bold":
                    text = f"<strong>{text}</strong>"
                elif mark["type"] == "italic":
                    text = f"<em>{text}</em>"
                elif mark["type"] == "underline":
                    text = f"<u>{text}</u>"
                elif mark["type"] == "strike":
                    text = f"<s>{text}</s>"
                elif mark["type"] == "code":
                    text = f"<code>{text}</code>"
                elif mark["type"] == "link":
                    href = sanitize_href(mark.get("attrs", {}).get("href", ""))
                    text = f'<a href="{_html_escape(href)}">{text}</a>'
            return text
        elif node_type == "bulletList":
            return f"<ul>{children}</ul>"
        elif node_type == "orderedList":
            return f"<ol>{children}</ol>"
        elif node_type == "listItem":
            return f"<li>{children}</li>"
        elif node_type == "blockquote":
            return f"<blockquote>{children}</blockquote>"
        elif node_type == "codeBlock":
            lang = _html_escape(node.get("attrs", {}).get("language", ""))
            code = (node.get("content", [{}])[0].get("text", "") if node.get("content") else "")
            code = code.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            return f'<pre><code class="language-{lang}">{code}</code></pre>'
        elif node_type == "hardBreak":
            return "<br>"
        elif node_type == "horizontalRule":
            return "<hr>"
        elif node_type == "image":
            src = _html_escape(node.get("attrs", {}).get("src", ""))
            alt = _html_escape(node.get("attrs", {}).get("alt", ""))
            return f'<img src="{src}" alt="{alt}">'
        elif node_type == "taskList":
            return f"<ul>{children}</ul>"
        elif node_type == "taskItem":
            checked = node.get("attrs", {}).get("checked", False)
            mark = "☑" if checked else "☐"
            return f"<li>{mark} {children}</li>"
        else:
            return children

    return render_node(content_json)


def html_to_markdown(html: str) -> str:
    """Simple HTML to Markdown converter for export."""
    md = html
    md = re.sub(r"<h1>(.*?)</h1>", r"# \1\n", md)
    md = re.sub(r"<h2>(.*?)</h2>", r"## \1\n", md)
    md = re.sub(r"<h3>(.*?)</h3>", r"### \1\n", md)
    md = re.sub(r"<h4>(.*?)</h4>", r"#### \1\n", md)
    md = re.sub(r"<h5>(.*?)</h5>", r"##### \1\n", md)
    md = re.sub(r"<h6>(.*?)</h6>", r"###### \1\n", md)
    md = re.sub(r"<strong>(.*?)</strong>", r"**\1**", md)
    md = re.sub(r"<em>(.*?)</em>", r"*\1*", md)
    md = re.sub(r"<u>(.*?)</u>", r"\1", md)
    md = re.sub(r"<s>(.*?)</s>", r"~~\1~~", md)
    md = re.sub(r"<code>(.*?)</code>", r"`\1`", md)
    md = re.sub(r'<a href="(.*?)">(.*?)</a>', r"[\2](\1)", md)
    md = re.sub(r"<blockquote>(.*?)</blockquote>", lambda m: "\n".join("> " + l for l in m.group(1).split("\n")) + "\n", md, flags=re.DOTALL)
    md = re.sub(r"<p>(.*?)</p>", r"\1\n\n", md, flags=re.DOTALL)
    md = re.sub(r"<br\s*/?>", "\n", md)
    md = re.sub(r"<hr\s*/?>", "\n---\n", md)
    md = re.sub(r"<li>(.*?)</li>", r"- \1\n", md, flags=re.DOTALL)
    md = re.sub(r"<ul>(.*?)</ul>", r"\1\n", md, flags=re.DOTALL)
    md = re.sub(r"<ol>(.*?)</ol>", r"\1\n", md, flags=re.DOTALL)
    md = re.sub(r"<pre><code[^>]*>(.*?)</code></pre>", lambda m: "```\n" + m.group(1) + "\n```\n", md, flags=re.DOTALL)
    # Strip remaining tags
    md = re.sub(r"<[^>]+>", "", md)
    # Clean up excessive newlines
    md = re.sub(r"\n{3,}", "\n\n", md)
    return md.strip()


def _latex_escape(text: str) -> str:
    """Escape special LaTeX characters."""
    for char, replacement in [
        ("\\", r"\textbackslash{}"),
        ("&", r"\&"), ("%", r"\%"), ("$", r"\$"),
        ("#", r"\#"), ("_", r"\_"), ("{", r"\{"),
        ("}", r"\}"), ("~", r"\textasciitilde{}"),
        ("^", r"\textasciicircum{}"),
    ]:
        text = text.replace(char, replacement)
    return text


def _render_latex_node(node: dict) -> str:
    """Recursively render a TipTap JSON node to LaTeX."""
    node_type = node.get("type", "")
    children = "".join(_render_latex_node(c) for c in node.get("content", []))

    if node_type == "doc":
        return children

    elif node_type == "paragraph":
        if not children.strip():
            return "\n"
        return children + "\n\n"

    elif node_type == "heading":
        level = node.get("attrs", {}).get("level", 1)
        cmds = {1: r"\section", 2: r"\subsection", 3: r"\subsubsection", 4: r"\paragraph", 5: r"\subparagraph"}
        cmd = cmds.get(level, r"\paragraph")
        return f"{cmd}{{{children.strip()}}}\n\n"

    elif node_type == "text":
        text = _latex_escape(node.get("text") or "")
        for mark in node.get("marks", []):
            mt = mark["type"]
            if mt == "bold":
                text = rf"\textbf{{{text}}}"
            elif mt == "italic":
                text = rf"\textit{{{text}}}"
            elif mt == "underline":
                text = rf"\uline{{{text}}}"
            elif mt == "strike":
                text = rf"\sout{{{text}}}"
            elif mt == "code":
                text = rf"\texttt{{{text}}}"
            elif mt == "link":
                href = mark.get("attrs", {}).get("href", "")
                text = rf"\href{{{href}}}{{{text}}}"
        return text

    elif node_type == "bulletList":
        items = []
        for child in node.get("content", []):
            item_text = "".join(_render_latex_node(c) for c in child.get("content", [])).strip()
            items.append(rf"\item {item_text}")
        return "\\begin{itemize}\n" + "\n".join(items) + "\n\\end{itemize}\n\n"

    elif node_type == "orderedList":
        items = []
        for child in node.get("content", []):
            item_text = "".join(_render_latex_node(c) for c in child.get("content", [])).strip()
            items.append(rf"\item {item_text}")
        return "\\begin{enumerate}\n" + "\n".join(items) + "\n\\end{enumerate}\n\n"

    elif node_type == "listItem":
        return children

    elif node_type == "blockquote":
        lines = children.strip().split("\n")
        quoted = "\n".join(rf"\quad \textit{{{l}}}" for l in lines if l.strip())
        return "\\begin{quote}\n" + quoted + "\n\\end{quote}\n\n"

    elif node_type == "codeBlock":
        lang = node.get("attrs", {}).get("language", "")
        code = ""
        if node.get("content"):
            code = node["content"][0].get("text", "")
        if lang:
            return f"\\begin{{lstlisting}}[language={lang}]\n{code}\n\\end{{lstlisting}}\n\n"
        return f"\\begin{{lstlisting}}\n{code}\n\\end{{lstlisting}}\n\n"

    elif node_type == "hardBreak":
        return "\\\\\n"

    elif node_type == "horizontalRule":
        return "\\noindent\\rule{\\textwidth}{0.4pt}\n\n"

    elif node_type == "taskList":
        items = []
        for child in node.get("content", []):
            checked = child.get("attrs", {}).get("checked", False)
            box = r"\checkbox" if checked else r"\square"
            item_text = "".join(_render_latex_node(c) for c in child.get("content", [])).strip()
            items.append(rf"\item[{box}] {item_text}")
        return "\\begin{itemize}\n" + "\n".join(items) + "\n\\end{itemize}\n\n"

    elif node_type == "taskItem":
        return children

    elif node_type == "image":
        src = node.get("attrs", {}).get("src", "")
        alt = node.get("attrs", {}).get("alt", "")
        return f"% [Image: {alt} — {src}]\n\n"

    return children


def tiptap_to_latex(content_json: dict | None, title: str = "") -> str:
    """Convert TipTap JSON document to a complete LaTeX source."""
    if not content_json or content_json.get("type") != "doc":
        body = ""
    else:
        body = _render_latex_node(content_json)

    escaped_title = _latex_escape(title) if title else "Untitled"

    return (
        r"\documentclass[12pt,a4paper]{article}" + "\n"
        r"\usepackage{xeCJK}" + "\n"
        r"\setCJKmainfont{Noto Sans CJK SC}" + "\n"
        r"\usepackage{hyperref}" + "\n"
        r"\usepackage{enumitem}" + "\n"
        r"\usepackage{listings}" + "\n"
        r"\usepackage{xcolor}" + "\n"
        r"\usepackage{ulem}" + "\n"
        r"\usepackage[margin=2.5cm]{geometry}" + "\n"
        r"\lstset{basicstyle=\ttfamily\small, breaklines=true, frame=single}" + "\n"
        r"\hypersetup{colorlinks=true, linkcolor=blue, urlcolor=blue}" + "\n\n"
        r"\title{" + escaped_title + "}\n"
        r"\date{\today}" + "\n\n"
        r"\begin{document}" + "\n"
        r"\maketitle" + "\n\n"
        + body + "\n"
        r"\end{document}" + "\n"
    )


def extract_plain_text(content_json) -> str:
    """Extract plain text from TipTap JSON for full-text search."""
    if content_json is None:
        return ""
    # Handle old double-serialized string format
    if isinstance(content_json, str):
        try:
            content_json = json.loads(content_json)
        except (json.JSONDecodeError, TypeError):
            return content_json  # Just return raw string if it's not valid JSON
    if not isinstance(content_json, dict) or content_json.get("type") != "doc":
        return ""

    parts: list[str] = []

    def walk(node: dict):
        if node.get("type") == "text":
            parts.append(node.get("text", ""))
        for child in node.get("content", []):
            walk(child)

    walk(content_json)
    return " ".join(parts)


def tiptap_to_docx(content_json: dict | None, title: str):
    """Convert TipTap JSON to a python-docx Document object."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDocument()

    # Title
    if title:
        doc.add_heading(title, level=0)

    if not content_json or content_json.get("type") != "doc":
        return doc

    def render_node(node: dict, list_style: str | None = None):
        node_type = node.get("type", "")

        if node_type == "doc":
            for child in node.get("content", []):
                render_node(child)

        elif node_type == "paragraph":
            p = doc.add_paragraph()
            for child in node.get("content", []):
                _add_inline(p, child)

        elif node_type == "heading":
            level = min(node.get("attrs", {}).get("level", 1), 4)
            p = doc.add_heading(level=level)
            for child in node.get("content", []):
                _add_inline(p, child)

        elif node_type == "bulletList":
            for child in node.get("content", []):
                render_node(child, list_style="List Bullet")

        elif node_type == "orderedList":
            for child in node.get("content", []):
                render_node(child, list_style="List Number")

        elif node_type == "listItem":
            p = doc.add_paragraph(style=list_style or "List Bullet")
            for child in node.get("content", []):
                if child.get("type") == "paragraph":
                    for inline in child.get("content", []):
                        _add_inline(p, inline)
                else:
                    _add_inline(p, child)

        elif node_type == "blockquote":
            for child in node.get("content", []):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                for inline in child.get("content", []):
                    _add_inline(p, inline)

        elif node_type == "codeBlock":
            code = ""
            if node.get("content"):
                code = node["content"][0].get("text", "")
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(code)
            run.font.name = "Courier New"
            run.font.size = Pt(9)

        elif node_type == "horizontalRule":
            doc.add_paragraph("─" * 60)

        elif node_type == "taskList":
            for child in node.get("content", []):
                checked = child.get("attrs", {}).get("checked", False)
                prefix = "☑ " if checked else "☐ "
                p = doc.add_paragraph()
                p.add_run(prefix)
                for sub in child.get("content", []):
                    if sub.get("type") == "paragraph":
                        for inline in sub.get("content", []):
                            _add_inline(p, inline)

        elif node_type == "image":
            alt = node.get("attrs", {}).get("alt", "Image")
            doc.add_paragraph(f"[{alt}]")

    def _add_inline(paragraph, node: dict):
        """Add inline content (text with marks) to a paragraph."""
        if node.get("type") != "text":
            # Recurse for nested inline nodes
            for child in node.get("content", []):
                _add_inline(paragraph, child)
            return

        text = node.get("text", "")
        marks = node.get("marks", [])
        run = paragraph.add_run(text)

        for mark in marks:
            mt = mark["type"]
            if mt == "bold":
                run.bold = True
            elif mt == "italic":
                run.italic = True
            elif mt == "underline":
                run.underline = True
            elif mt == "strike":
                run.font.strike = True
            elif mt == "code":
                run.font.name = "Courier New"
                run.font.size = Pt(9)

    render_node(content_json)
    return doc


def _parse_inline(text: str) -> list[dict]:
    """Parse inline Markdown formatting into TipTap text nodes."""
    nodes: list[dict] = []
    # Pattern to match **bold**, *italic*, `code`, [link](url)
    pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[(.+?)\]\((.+?)\))"
    last_end = 0

    for m in re.finditer(pattern, text):
        # Add preceding plain text
        if m.start() > last_end:
            nodes.append({"type": "text", "text": text[last_end:m.start()]})

        if m.group(2):  # **bold**
            nodes.append({"type": "text", "text": m.group(2), "marks": [{"type": "bold"}]})
        elif m.group(3):  # *italic*
            nodes.append({"type": "text", "text": m.group(3), "marks": [{"type": "italic"}]})
        elif m.group(4):  # `code`
            nodes.append({"type": "text", "text": m.group(4), "marks": [{"type": "code"}]})
        elif m.group(5) and m.group(6):  # [text](url)
            nodes.append({"type": "text", "text": m.group(5), "marks": [{"type": "link", "attrs": {"href": m.group(6)}}]})

        last_end = m.end()

    # Remaining text
    if last_end < len(text):
        nodes.append({"type": "text", "text": text[last_end:]})

    if not nodes:
        nodes.append({"type": "text", "text": text})

    return nodes


def markdown_to_tiptap(md_text: str) -> dict:
    """Convert Markdown text to TipTap JSON format."""
    lines = md_text.split("\n")
    content: list[dict] = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            content.append({
                "type": "heading",
                "attrs": {"level": min(level, 6)},
                "content": _parse_inline(text),
            })
            i += 1
            continue

        # Code blocks
        if line.strip().startswith("```"):
            lang = line.strip()[3:].strip()
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            code_text = "\n".join(code_lines)
            content.append({
                "type": "codeBlock",
                "attrs": {"language": lang} if lang else {},
                "content": [{"type": "text", "text": code_text}] if code_text else [],
            })
            continue

        # Blockquote
        if line.startswith(">"):
            text = re.sub(r"^>\s?", "", line)
            content.append({
                "type": "blockquote",
                "content": [{"type": "paragraph", "content": _parse_inline(text)}],
            })
            i += 1
            continue

        # Bullet list
        if re.match(r"^[\-\*]\s+", line):
            items: list[dict] = []
            while i < len(lines) and re.match(r"^[\-\*]\s+", lines[i]):
                text = re.sub(r"^[\-\*]\s+", "", lines[i])
                items.append({
                    "type": "listItem",
                    "content": [{"type": "paragraph", "content": _parse_inline(text)}],
                })
                i += 1
            content.append({"type": "bulletList", "content": items})
            continue

        # Ordered list
        if re.match(r"^\d+\.\s+", line):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i]):
                text = re.sub(r"^\d+\.\s+", "", lines[i])
                items.append({
                    "type": "listItem",
                    "content": [{"type": "paragraph", "content": _parse_inline(text)}],
                })
                i += 1
            content.append({"type": "orderedList", "content": items})
            continue

        # Horizontal rule
        if re.match(r"^(\-{3,}|\*{3,}|_{3,})\s*$", line):
            content.append({"type": "horizontalRule"})
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        content.append({
            "type": "paragraph",
            "content": _parse_inline(line),
        })
        i += 1

    return {"type": "doc", "content": content}


# ── DOCX to TipTap Conversion ──────────────────────────────────────


def _extract_runs(para) -> list[dict]:
    """Extract text runs from a DOCX paragraph with mark preservation."""
    nodes = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        marks = []
        if run.bold:
            marks.append({"type": "bold"})
        if run.italic:
            marks.append({"type": "italic"})
        if run.underline:
            marks.append({"type": "underline"})
        if run.font.strike:
            marks.append({"type": "strike"})
        if run.font.superscript:
            marks.append({"type": "superscript"})
        if run.font.subscript:
            marks.append({"type": "subscript"})
        node = {"type": "text", "text": text}
        if marks:
            node["marks"] = marks
        nodes.append(node)
    return nodes or [{"type": "text", "text": para.text or ""}]


def _get_list_type(doc, num_id: str) -> str:
    """Determine if a list is ordered or bullet by checking numbering XML."""
    try:
        from docx.oxml.ns import qn
        numbering_part = doc.part.numbering_part
        numbering_elem = numbering_part.numbering_definitions._numbering
        # Find the abstractNumId for this numId
        num_elems = numbering_elem.findall(qn("w:num"))
        for num_elem in num_elems:
            if num_elem.get(qn("w:numId")) == str(num_id):
                abstract_num_id_elem = num_elem.find(qn("w:abstractNumId"))
                if abstract_num_id_elem is not None:
                    abstract_num_id = abstract_num_id_elem.get(qn("w:val"))
                    # Find the abstractNum and its first level's numFmt
                    abstract_nums = numbering_elem.findall(qn("w:abstractNum"))
                    for anum in abstract_nums:
                        if anum.get(qn("w:abstractNumId")) == abstract_num_id:
                            lvl = anum.find(qn("w:lvl"))
                            if lvl is not None:
                                num_fmt = lvl.find(qn("w:numFmt"))
                                if num_fmt is not None:
                                    fmt_val = num_fmt.get(qn("w:val"))
                                    if fmt_val == "decimal":
                                        return "orderedList"
    except Exception:
        pass
    return "bulletList"


def docx_to_tiptap(file_bytes: bytes) -> dict:
    """Convert DOCX bytes to TipTap JSON, preserving headings, bold/italic, lists, tables."""
    import io
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(file_bytes))
    content_nodes: list[dict] = []
    pending_list: dict | None = None  # {type, items}

    def flush_list():
        nonlocal pending_list
        if pending_list and pending_list["items"]:
            content_nodes.append({
                "type": pending_list["type"],
                "content": pending_list["items"],
            })
        pending_list = None

    # Walk paragraphs
    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower()
        text = para.text.strip()

        # ── Heading detection ──
        if style_name.startswith("heading"):
            flush_list()
            try:
                level = int(style_name.replace("heading", "").strip())
            except ValueError:
                level = 1
            level = max(1, min(level, 6))
            content_nodes.append({
                "type": "heading",
                "attrs": {"level": level},
                "content": _extract_runs(para),
            })
            continue

        # ── List detection ──
        pPr = para._element.find(qn("w:pPr"))
        if pPr is not None:
            numPr = pPr.find(qn("w:numPr"))
            if numPr is not None:
                numId_el = numPr.find(qn("w:numId"))
                if numId_el is not None:
                    num_id = numId_el.get(qn("w:val"), "0")
                    if num_id != "0":
                        list_type = _get_list_type(doc, num_id)
                        list_item = {
                            "type": "listItem",
                            "content": [{"type": "paragraph", "content": _extract_runs(para)}],
                        }
                        if pending_list and pending_list["type"] == list_type:
                            pending_list["items"].append(list_item)
                        else:
                            flush_list()
                            pending_list = {"type": list_type, "items": [list_item]}
                        continue

        # ── Regular paragraph ──
        flush_list()
        if text:
            content_nodes.append({
                "type": "paragraph",
                "content": _extract_runs(para),
            })
        else:
            content_nodes.append({"type": "paragraph"})

    flush_list()

    # ── Tables ──
    for table in doc.tables:
        table_rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_content = []
                for cell_para in cell.paragraphs:
                    cell_content.append({
                        "type": "paragraph",
                        "content": _extract_runs(cell_para),
                    })
                cells.append({"type": "tableCell", "content": cell_content})
            table_rows.append({"type": "tableRow", "content": cells})
        content_nodes.append({"type": "table", "content": table_rows})

    # ── Images (inline in document order) ──
    try:
        import base64
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_bytes = rel.target_part.blob
                content_type = rel.target_part.content_type
                b64 = base64.b64encode(image_bytes).decode()
                src = f"data:{content_type};base64,{b64}"
                content_nodes.append({
                    "type": "image",
                    "attrs": {"src": src, "alt": "imported image"},
                })
    except Exception:
        pass  # Images are optional

    return {"type": "doc", "content": content_nodes}


# ── LaTeX to TipTap Conversion ─────────────────────────────────────


def latex_to_tiptap(text: str) -> dict:
    """Convert LaTeX source to TipTap JSON, preserving structure."""
    content_nodes: list[dict] = []
    pending_list: dict | None = None

    def flush_list():
        nonlocal pending_list
        if pending_list and pending_list["items"]:
            content_nodes.append({
                "type": pending_list["type"],
                "content": pending_list["items"],
            })
        pending_list = None

    def make_list_item(text: str) -> dict:
        return {
            "type": "listItem",
            "content": [{"type": "paragraph", "content": [{"type": "text", "text": text.strip()}]}],
        }

    # Remove preamble (everything before \begin{document})
    doc_match = re.search(r"\\begin\{document\}", text)
    if doc_match:
        text = text[doc_match.end():]
    # Remove \end{document} and after
    text = re.sub(r"\\end\{document\}.*", "", text, flags=re.DOTALL)

    # Remove comments
    text = re.sub(r"%.*$", "", text, flags=re.MULTILINE)

    # Remove common preamble commands
    text = re.sub(r"\\(?:usepackage|documentclass|newcommand|renewcommand|setlength|geometry|bibliography|bibliographystyle|maketitle|tableofcontents|label|ref|cite|pageref|eqref)(?:\[[^\]]*\])*\{[^}]*\}", "", text)

    # Split into lines for processing
    lines = text.split("\n")
    i = 0
    in_list = None  # "bulletList" or "orderedList"

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            flush_list()
            i += 1
            continue

        # ── Section headings ──
        heading_match = re.match(r"\\(chapter|part)\{([^}]+)\}", line)
        if heading_match:
            flush_list()
            content_nodes.append({
                "type": "heading",
                "attrs": {"level": 1},
                "content": [{"type": "text", "text": heading_match.group(2).strip()}],
            })
            i += 1
            continue

        heading_match = re.match(r"\\(section)\{([^}]+)\}", line)
        if heading_match:
            flush_list()
            content_nodes.append({
                "type": "heading",
                "attrs": {"level": 1},
                "content": [{"type": "text", "text": heading_match.group(2).strip()}],
            })
            i += 1
            continue

        heading_match = re.match(r"\\(subsection)\{([^}]+)\}", line)
        if heading_match:
            flush_list()
            content_nodes.append({
                "type": "heading",
                "attrs": {"level": 2},
                "content": [{"type": "text", "text": heading_match.group(2).strip()}],
            })
            i += 1
            continue

        heading_match = re.match(r"\\(subsubsection)\{([^}]+)\}", line)
        if heading_match:
            flush_list()
            content_nodes.append({
                "type": "heading",
                "attrs": {"level": 3},
                "content": [{"type": "text", "text": heading_match.group(2).strip()}],
            })
            i += 1
            continue

        # ── List items ──
        if re.match(r"\\item\s+", line):
            item_text = re.sub(r"^\\item\s+", "", line).strip()
            item_text = _clean_latex_inline(item_text)
            list_type = "orderedList" if in_list == "orderedList" else "bulletList"

            if pending_list and pending_list["type"] == list_type:
                pending_list["items"].append(make_list_item(item_text))
            else:
                flush_list()
                in_list = "bulletList"
                pending_list = {"type": "bulletList", "items": [make_list_item(item_text)]}
            i += 1
            continue

        # ── List environments ──
        if re.match(r"\\begin\{(itemize|enumerate|description)\}", line):
            flush_list()
            env_match = re.match(r"\\begin\{(itemize|enumerate|description)\}", line)
            env_name = env_match.group(1)
            in_list = "orderedList" if env_name == "enumerate" else "bulletList"
            pending_list = {"type": in_list, "items": []}
            i += 1
            continue

        if re.match(r"\\end\{(itemize|enumerate|description)\}", line):
            flush_list()
            in_list = None
            i += 1
            continue

        # ── Blockquote ──
        if re.match(r"\\begin\{quote\}", line):
            flush_list()
            quote_lines = []
            i += 1
            while i < len(lines) and not re.match(r"\\end\{quote\}", lines[i]):
                if lines[i].strip():
                    quote_lines.append(lines[i].strip())
                i += 1
            if quote_lines:
                content_nodes.append({
                    "type": "blockquote",
                    "content": [{"type": "paragraph", "content": [{"type": "text", "text": " ".join(quote_lines)}]}],
                })
            i += 1
            continue

        # ── Code block ──
        if re.match(r"\\begin\{(verbatim|lstlisting|minted)\}", line):
            flush_list()
            code_lines = []
            i += 1
            while i < len(lines) and not re.match(r"\\end\{(verbatim|lstlisting|minted)\}", lines[i]):
                code_lines.append(lines[i])
                i += 1
            content_nodes.append({
                "type": "codeBlock",
                "content": [{"type": "text", "text": "\n".join(code_lines)}],
            })
            i += 1
            continue

        # ── Horizontal rule ──
        if re.match(r"\\(hrule|hr|rule)", line):
            flush_list()
            content_nodes.append({"type": "horizontalRule"})
            i += 1
            continue

        # ── Skip unknown environments ──
        if re.match(r"\\begin\{", line):
            i += 1
            continue
        if re.match(r"\\end\{", line):
            i += 1
            continue

        # ── Regular paragraph ──
        flush_list()
        cleaned = _clean_latex_inline(line)
        if cleaned:
            content_nodes.append({
                "type": "paragraph",
                "content": [{"type": "text", "text": cleaned}],
            })
        i += 1

    flush_list()
    return {"type": "doc", "content": content_nodes if content_nodes else [{"type": "paragraph"}]}


def _clean_latex_inline(text: str) -> str:
    """Clean LaTeX inline commands, preserving content."""
    # Bold
    text = re.sub(r"\\textbf\{([^}]*)\}", r"\1", text)
    # Italic
    text = re.sub(r"\\textit\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\emph\{([^}]*)\}", r"\1", text)
    # Underline
    text = re.sub(r"\\underline\{([^}]*)\}", r"\1", text)
    # Superscript/subscript
    text = re.sub(r"\^{([^}]*)}", r"^\1", text)
    text = re.sub(r"_{([^}]*)}", r"_\1", text)
    # Remove remaining commands with braces
    text = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])*\{([^}]*)\}", r"\1", text)
    # Remove remaining backslash commands without braces
    text = re.sub(r"\\[a-zA-Z]+\*?", "", text)
    # Remove braces
    text = text.replace("{", "").replace("}", "")
    # Math delimiters
    text = re.sub(r"\$\$(.*?)\$\$", r" \1 ", text, flags=re.DOTALL)
    text = re.sub(r"\$(.*?)\$", r"\1", text, flags=re.DOTALL)
    return text.strip()
